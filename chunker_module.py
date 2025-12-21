import os
import pandas as pd
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader, UnstructuredPowerPointLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

def detect_permission_from_content(text_content):
    if not text_content:
        return "user"
        
    text_lower = text_content.lower()[:5000]

    # 1. Seviye: Admin
    admin_keywords = ["gizli", "confidential", "passwords", "şifreler", "yönetim kurulu", "admin only"]
    if any(k in text_lower for k in admin_keywords):
        return "admin"

    # 2. Seviye: Manager
    manager_keywords = ["maliyet", "bütçe", "budget", "finans", "salary", "maaş", "forecast"]
    if any(k in text_lower for k in manager_keywords):
        return "manager"

    # 3. Seviye: Editor
    editor_keywords = ["taslak", "draft", "düzenlenecek", "teknik şartname"]
    if any(k in text_lower for k in editor_keywords):
        return "editor"

    # Varsayılan: User
    return "user"

class DriveDocumentProcessor:
    def __init__(self, chunk_size=1000, chunk_overlap=100):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        self.processed_chunks = []

    def process_excel(self, file_path, fixed_permission=None):
        try:
            xls = pd.read_excel(file_path, sheet_name=None)
            filename = os.path.basename(file_path)

            for sheet_name, df in xls.items():
                df.dropna(how='all', axis=0, inplace=True)
                df.dropna(how='all', axis=1, inplace=True)

                for col in df.columns:
                    if pd.api.types.is_datetime64_any_dtype(df[col]):
                        df[col] = df[col].dt.strftime('%Y-%m-%d')

                try:
                    # Markdown tablosu oluştur
                    markdown_table = df.to_markdown(index=False)
                except ImportError:
                    markdown_table = df.to_string(index=False)

                permission = fixed_permission if fixed_permission else detect_permission_from_content(markdown_table)

                table_text = (
                    f"KAYNAK: {filename} > {sheet_name}\n"
                    f"TÜR: Excel Tablosu\n"
                    f"İÇERİK:\n"
                    f"{markdown_table}"
                )

                metadata = {
                    "source": filename,
                    "sheet": sheet_name,
                    "file_type": "excel",
                    "permission": permission
                }

                self.processed_chunks.append(Document(page_content=table_text, metadata=metadata))
            
            return True, f"Excel processed: {len(xls)} sheets"
        except Exception as e:
            return False, str(e)

    def process_word(self, file_path, fixed_permission=None):
        try:
            doc = DocxDocument(file_path)
            filename = os.path.basename(file_path)

            full_text = "\n".join([p.text for p in doc.paragraphs])
            permission = fixed_permission if fixed_permission else detect_permission_from_content(full_text)

            # Metadata'yı her chunk için güncellemek gerek
            text_chunks = self.text_splitter.create_documents([full_text])
            for chunk in text_chunks:
                chunk.metadata.update({
                    "source": filename,
                    "file_type": "word",
                    "permission": permission
                })
                self.processed_chunks.append(chunk)
            
            # Tabloları işle
            for i, table in enumerate(doc.tables):
                 if not table.rows: continue
                 
                 # Basit tablo işleme
                 rows_text = []
                 for row in table.rows:
                     rows_text.append(" | ".join([cell.text.strip() for cell in row.cells]))
                 
                 table_content = "\n".join(rows_text)
                 
                 chunk = Document(
                     page_content=f"TABLO:\n{table_content}",
                     metadata={
                        "source": filename,
                        "file_type": "word_table",
                        "permission": permission
                     }
                 )
                 self.processed_chunks.append(chunk)

            return True, "Word processed"
        except Exception as e:
            return False, str(e)

    def process_pdf_ppt(self, file_path, file_type, fixed_permission=None):
        try:
            if file_type == "pdf":
                loader = PyPDFLoader(file_path)
            else:
                loader = UnstructuredPowerPointLoader(file_path)
            
            raw_docs = loader.load()
            
            full_text = " ".join([d.page_content for d in raw_docs])
            permission = fixed_permission if fixed_permission else detect_permission_from_content(full_text)
            
            chunks = self.text_splitter.split_documents(raw_docs)
            filename = os.path.basename(file_path)
            
            for chunk in chunks:
                chunk.metadata.update({
                    "source": filename,
                    "permission": permission,
                    "file_type": file_type
                })
                self.processed_chunks.append(chunk)
                
            return True, f"{file_type.upper()} processed"
        except Exception as e:
            return False, str(e)

def process_file_wrapper(file_path, fixed_permission=None):
    """
    Main entry point for API.
    Determines file type and calls appropriate processor.
    Returns: (list_of_documents, error_message)
    """
    processor = DriveDocumentProcessor()
    filename = os.path.basename(file_path).lower()
    
    success = False
    msg = ""

    if filename.endswith(".xlsx") or filename.endswith(".xls"):
        success, msg = processor.process_excel(file_path, fixed_permission)
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        success, msg = processor.process_word(file_path, fixed_permission)
    elif filename.endswith(".pdf"):
        success, msg = processor.process_pdf_ppt(file_path, "pdf", fixed_permission)
    elif filename.endswith(".pptx") or filename.endswith(".ppt"):
        success, msg = processor.process_pdf_ppt(file_path, "powerpoint", fixed_permission)
    else:
        return [], f"Unsupported file type: {filename}"

    if success:
        return processor.processed_chunks, None
    else:
        return [], msg
